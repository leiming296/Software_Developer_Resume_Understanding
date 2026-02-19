"""
Word document parser implementation.
This module provides a concrete implementation of FileParser for Word documents
using the python-docx library.
"""

from docx import Document
from pathlib import Path
from .base import FileParser


class WordParser(FileParser):
    """
    Concrete implementation of FileParser for Word documents (.docx).
    Uses python-docx library to extract text from Word documents.
    Preserves document structure by extracting text from all paragraphs.
    """

    def parse(self, file_path: str) -> str:
        """
        Parse a Word document and extract its text content.
        Args:file_path: Path to the Word document (.docx)
        Returns:Extracted text content from all paragraphs
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Word document not found: {file_path}")

        if not path.suffix.lower() in ['.docx', '.doc']:
            raise ValueError(f"File is not a Word document: {file_path}")

        try:
            doc = Document(file_path)

            # Extract text from all paragraphs
            text_content = []

            # Extract from headers (often contains name in resumes)
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)

            # Extract from main document paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Join all paragraphs with newlines
            full_text = '\n'.join(text_content)

            if not full_text.strip():
                raise ValueError("Word document appears to be empty")

            return full_text

        except Exception as e:
            raise Exception(f"Failed to parse Word document: {str(e)}")
